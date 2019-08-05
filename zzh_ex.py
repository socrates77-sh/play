# history:
# 2019/08/05  v1.0  initial

# import docx
import msvcrt
import os
import random
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


VERSION = '1.0'

TITLE = '张芷蘅的乘法练习'
PAGE_MARGIN = 0.5


def page_init(doc):
    distance = Inches(PAGE_MARGIN)
    sec = doc.sections[0]
    sec.left_margin = distance
    sec.right_margin = distance
    sec.top_margin = distance
    sec.bottom_margin = distance


def change_font(obj, fontname='微软雅黑', size=None):
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size


def title(doc, text):
    heading = doc.add_heading(text, 1)
    font = heading.style.font
    font.size = Pt(24)
    font.bold = False
    font.underline = True
    font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format = heading.style.paragraph_format
    paragraph_format.space_after = Pt(20)


def gen_a_mul():
    m1 = random.randint(1, 9)
    m2 = random.randint(1, 9)
    return '%dx%d=' % (m1, m2)


def print_a_line(doc):
    s1 = gen_a_mul()
    s2 = gen_a_mul()
    s3 = gen_a_mul()
    s4 = gen_a_mul()
    p = doc.add_paragraph('%-10s%-10s%-10s%-10s' % (s1, s2, s3, s4))
    # p.space_before(Pt(50))


def print_version(version):
    print('=' * 70)
    print('%s version: %s' % (os.path.basename(__file__), VERSION))
    print('=' * 70)


def wait_any_key():
    print('press any key to exit...')
    msvcrt.getch()


def main():
    print_version(VERSION)
    file_name = TITLE + '.docx'

    doc = Document()
    page_init(doc)
    title(doc, TITLE)

    change_font(doc.styles['Normal'], fontname='Consolas', size=Pt(24))

    for i in range(15):
        print_a_line(doc)

    try:
        doc.save(file_name)
        print('[Save] %s' % file_name)
    except PermissionError:
        print('请先关闭文件: %s' % file_name)

    wait_any_key()


if __name__ == '__main__':
    main()
